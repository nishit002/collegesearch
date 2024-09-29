import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from docx.shared import Pt

# Function to create a table in Word document
def add_table_to_doc(doc, data, headers):
    table = doc.add_table(rows=1, cols=len(headers))
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header

    for _, row in data.iterrows():
        row_cells = table.add_row().cells
        for i, header in enumerate(headers):
            row_cells[i].text = str(row[header])

# Function to create Word document with the selected college data
def create_word_document(college_info, ranking_data, placement_data, awards_data, faculty_data, recruiters_data, course_data, admission_data, contact_data, facilities_data, scholarships_data, cutoff_data, affiliation_data, approval_data):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Title
    doc.add_heading(f"{college_info['college_name']} Information", level=1).style = style

    # General Info
    doc.add_paragraph(f"{college_info['college_name']} was established in the {college_info['establishment_year']} and is located in {college_info['city']}, {college_info['state']}.", style=style)
    doc.add_paragraph(f"The college is known for its {college_info['usp']}. It is a Coed College{'.' if college_info['is_coed'] == 'Yes' else ' and is not a Coed College.'}", style=style)
    nirf_rank = college_info['nirf_rank'] if pd.notna(college_info['nirf_rank']) else 'N/A'  # Handle potential missing NIRF rank
    doc.add_paragraph(f"The NIRF rank of the college is {nirf_rank}.", style=style)

    # Approvals and Affiliations
    approved_by = ", ".join(approval_data['approval_body'].astype(str).tolist())
    affiliated_to = ", ".join(affiliation_data['affiliated_university'].astype(str).tolist())
    doc.add_paragraph(f"It has been approved by {approved_by}.", style=style)
    doc.add_paragraph(f"This college has a wide range of courses like {', '.join(course_data['course_name'].astype(str).tolist())}. It is affiliated with {affiliated_to}.", style=style)

    # NAAC Ranking (add a NAAC column to your College data)
    doc.add_paragraph(f"NAAC has ranked the college at {college_info.get('naac_rank', 'N/A')}.", style=style)  # Handle potential missing NAAC rank

    # Rankings Table
    doc.add_heading('Rankings', level=2).style = style
    if not ranking_data.empty:
        add_table_to_doc(doc, ranking_data, ['ranking_body', 'rank'])

    # Placements 
    doc.add_heading(f"{college_info['college_name']} Placements", level=2).style = style
    if not placement_data.empty:
        for index, row in placement_data.iterrows():
            doc.add_paragraph(f"The college placement records for {row['course_name']} are as follows, the Highest Package is INR {row['highest_package']} and the Average Package is INR {row['average_package']}.", style=style)

    # Top Recruiters and Industries
    doc.add_paragraph(f"The Top Recruiters of the {college_info['college_name']} are {', '.join(recruiters_data['recruiter_name'].astype(str).tolist())}.", style=style)  # Assuming you have a dedicated "Industries" column in your recruiters data
    doc.add_paragraph(f"The Top Industries that the students got hired from are  {', '.join(recruiters_data['recruiter_name'].astype(str).tolist())}.", style=style)  # Placeholder - replace with actual industry data

    # Awards
    doc.add_heading(f"{college_info['college_name']} Awards", level=2).style = style
    if not awards_data.empty:
        for index, row in awards_data.iterrows():
            doc.add_paragraph(f"The college has been awarded with awards like {row['award_name']} by {row['awarding_body']} in the year {row['year']}.", style=style)
    # Alumni (add an alumni column to your Awards data or a new sheet)
    doc.add_paragraph(f"The college has some famous alumni from a variety of fields, some of the Top Alumni are {', '.join(awards_data['award_name'].astype(str).tolist())}", style=style)  # Placeholder - replace with actual alumni data

    # Scholarships
    doc.add_heading(f"{college_info['college_name']} Scholarships", level=2).style = style
    if not scholarships_data.empty:
        for index, row in scholarships_data.iterrows():
            doc.add_paragraph(f"Name: {row['scholarship_name']}", style=style)
            doc.add_paragraph(f"Description: {row['description']}", style=style)

    # Faculty Table
    doc.add_heading(f"{college_info['college_name']} Faculty", level=2).style = style
    if not faculty_data.empty:
        add_table_to_doc(doc, faculty_data, ['faculty_name', 'position', 'specialty', 'education'])

    # Contact and Infrastructure
    # ... (other parts of the function) ...

    # Contact and Infrastructure
    doc.add_heading(f"{college_info['college_name']} Address", level=2).style = style
    
    # Access contact details from contact_data
    contact_info = contact_data[contact_data['college_id'] == college_info['college_id']].iloc[0]
    
    doc.add_paragraph(f"The college is located in {contact_info['address']}. You can contact the college by reaching out to the phone number: {contact_info['phone_number']} Email: {contact_info['email']}", style=style)


    # Courses and Fees
    doc.add_heading("Courses and Fees", level=2).style = style
    total_courses = len(course_data)
    doc.add_paragraph(f"The {college_info['college_name']} Provides courses in undergraduate, Postgraduate, and Doctoral programs along with various vocational, technical, and online courses. In total, the college has {total_courses}.", style=style)

    for index, row in course_data.iterrows():
        doc.add_paragraph(f"It provides {row['course_name']}, which is a {row['level']} course and is of {row['duration']}, the fee for the course is {row['fee']}.", style=style)
        # Add placement details if available in placement_data
        placement_info = placement_data[placement_data['course_name'] == row['course_name']]
        if not placement_info.empty:
            doc.add_paragraph(f"The Placement for this course in the <year> was {placement_info['highest_package'].iloc[0]}, and {placement_info['average_package'].iloc[0]}.", style=style)  # Replace <year> with actual year
        # Add specialization details (assuming you have a "specializations" column in your course data)
        doc.add_paragraph(f"The specializations for this course include: {row['specialization']}.", style=style)  # Placeholder - replace with actual specialization data
        # Add information about course type, mode, and examination frequency (add these columns to your course data)
        doc.add_paragraph("The Course is a Full-time course. It is provided on campus. It is a degree course and takes examinations on a <semester/trimester/Yearly> basis.", style=style)  # Placeholder - replace with actual data

        # Add admission details if available in admission_data
        admission_info = admission_data[admission_data['course_name'] == row['course_name']]
        if not admission_info.empty:
            doc.add_paragraph(f"Admission for this course starts on {admission_info['start_date'].iloc[0].strftime('%Y-%m-%d')} and ends on {admission_info['end_date'].iloc[0].strftime('%Y-%m-%d')}.", style=style)
        doc.add_paragraph(f"The eligibility criteria for the course is {row['admission_criteria']}.", style=style)  # Assuming you have an "admission_criteria" column in your course data
        doc.add_paragraph(f"The total number of seats for this course is {row['seats']}.", style=style)

    # Admission Process
    doc.add_heading("Admission Process", level=2).style = style
    doc.add_paragraph("To apply for admission to this college, you can follow these steps:", style=style)
    doc.add_paragraph("1. Go to the <official website of college> and register with all your credentials.", style=style)  # Replace with actual website
    doc.add_paragraph("2. Log in to the admission portal.", style=style)
    doc.add_paragraph("3. Submit all of your documents.", style=style)
    doc.add_paragraph("4. Pay the required fee as per the portal.", style=style)
    doc.add_paragraph("5. Keep the printout of the Fees payment for future reference.", style=style)

    # Documents Required
    doc.add_heading("Documents Required for Admission", level=2).style = style
    doc.add_paragraph("Application form", style=style)
    doc.add_paragraph("Passport-sized photographs", style=style)
    # ... add other documents from your template ...

    # Admission Helpline 
    doc.add_paragraph(f"In case you have any queries you can contact the Admission Helpline Number: {contact_data['phone_number'].iloc[0]}.", style=style)

    # Cutoff Information (if available)
    doc.add_heading("College Cutoff", level=2).style = style
    if not cutoff_data.empty:
        doc.add_paragraph("The cutoff is the minimum eligibility required by the institute in order to enter students into various programs.", style=style)
        # You can add a table here to display cutoff_data if needed
        add_table_to_doc(doc, cutoff_data, ['course_name', 'cutoff_score'])
    else:
        doc.add_paragraph("Cutoff information is not available.", style=style)

    return doc

# Function to download Word file
def download_word_file(doc):
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
    # Add a file uploader so the user can upload the Excel file
st.title("College Information Portal")

uploaded_file = st.file_uploader("Upload your Excel file", type=["xlsx"])

if uploaded_file is not None:
    # Read the uploaded Excel file
    excel_data = pd.ExcelFile(uploaded_file)

    # Load data from the sheets into dataframes
    college_df = pd.read_excel(excel_data, 'College')
    ranking_df = pd.read_excel(excel_data, 'Ranking')
    placement_df = pd.read_excel(excel_data, 'Placement')
    course_df = pd.read_excel(excel_data, 'Courses')
    faculty_df = pd.read_excel(excel_data, 'Faculty')
    recruiters_df = pd.read_excel(excel_data, 'Recruiters')
    awards_df = pd.read_excel(excel_data, 'Awards')
    admission_df = pd.read_excel(excel_data, 'Admission')
    contact_df = pd.read_excel(excel_data, 'Contact_Details')
    facilities_df = pd.read_excel(excel_data, 'Facilities')
    scholarships_df = pd.read_excel(excel_data, 'Scholarship')
    cutoff_df = pd.read_excel(excel_data, 'Cutoff')
    affiliation_df = pd.read_excel(excel_data, 'Affiliation')
    approval_df = pd.read_excel(excel_data, 'Approval')

    # Sidebar for College Selection
    st.sidebar.header('Select College')
    college_name = st.sidebar.selectbox('College Name', college_df['college_name'].unique())

    # Display General Information about the selected college
    st.header(f'{college_name} Information')

    college_info = college_df[college_df['college_name'] == college_name].iloc[0]
    st.write(f"**{college_name}** was established in {college_info['establishment_year']} and is located in {college_info['city']}, {college_info['state']}.")
    st.write(f"The college is known for its {college_info['usp']}. It is a {'Coed' if college_info['is_coed'] == 'Yes' else 'Non-Coed'} college.")
    st.write(f"The NIRF rank of the college is {college_info['nirf_rank']}.")

    # ... (You can add more code here to display data on the Streamlit app if needed) ...
    # Generate and download Word document
    doc = create_word_document(college_info, ranking_df, placement_df, awards_df, 
                               faculty_df, recruiters_df, course_df, admission_df, 
                               contact_df, facilities_df, scholarships_df, cutoff_df,
                               affiliation_df, approval_df)
    buffer = download_word_file(doc)

    st.download_button(
        label="Download College Information in Word",
        data=buffer,
        file_name=f"{college_name}_info.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
else:
    st.write("Please upload the Excel file to get started.")
